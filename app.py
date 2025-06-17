import streamlit as st
import pandas as pd
from docx import Document

def extract_invoice_data(file):
    df = pd.read_excel(file, engine='xlrd')
    
    client_row = df[df.iloc[:, 0].astype(str).str.contains("Nom du client", na=False)]
    client = client_row.iloc[0, 1] if not client_row.empty else "Non trouvé"
    
    contract_row = df[df.iloc[:, 0].astype(str).str.contains("# Contrat", na=False)]
    contract = contract_row.iloc[0, 1] if not contract_row.empty else "Non trouvé"
    
    tec_row = df[df.iloc[:, 0].astype(str).str.contains("Solde TEC net", na=False)]
    tec_balance = tec_row.iloc[0, 1] if not tec_row.empty else "Non trouvé"
    
    last_invoice_index = df[df.iloc[:, 4].astype(str).str.contains("Facture standard", na=False)].index[-1]
    
    descriptions = []
    total_fees = 0
    for i in range(last_invoice_index + 1, len(df)):
        if pd.isna(df.iloc[i, 14]):
            break
        descriptions.append(str(df.iloc[i, 9]))
        total_fees += df.iloc[i, 14]
    
    descriptions = list(set(descriptions))
    descriptions = [desc.lower().capitalize() for desc in descriptions]
    
    return client, contract, tec_balance, descriptions, total_fees

def generate_invoice(client, contract, tec_balance, descriptions, total_fees):
    doc = Document()
    doc.add_heading('Brouillon de facture', level=1)
    doc.add_heading('Informations générales', level=2)
    doc.add_paragraph(f'Client : {client}')
    doc.add_paragraph(f'# Contrat : {contract}')
    doc.add_paragraph(f'Solde TEC net : {tec_balance}')
    doc.add_heading('Descriptions des travaux effectués', level=2)
    for desc in descriptions:
        doc.add_paragraph(f'- {desc}')
    doc.add_heading('Total des honoraires', level=2)
    doc.add_paragraph(f'55 % du total : {total_fees * 0.55:.2f} $ CAD')
    doc.add_paragraph(f'65 % du total : {total_fees * 0.65:.2f} $ CAD')
    doc.save('invoice_draft.docx')

st.title('Générateur de brouillon de facture')

uploaded_file = st.file_uploader("Téléverser un fichier Excel de type WIP", type=["xls"])

if uploaded_file is not None:
    client, contract, tec_balance, descriptions, total_fees = extract_invoice_data(uploaded_file)
    generate_invoice(client, contract, tec_balance, descriptions, total_fees)
    st.success('Le brouillon de facture a été généré avec succès.')
    with open('invoice_draft.docx', 'rb') as f:
        st.download_button('Télécharger le brouillon de facture', f, file_name='invoice_draft.docx')
